# import re
# import phonenumbers
# import spacy
# from pypdf import PdfReader
# from docx import Document

# class ResumeParser:
#     def __init__(self, file_path):
#         self.file_path = file_path
#         self.raw_text = self.extract_text(file_path)
#         self.text = self.clean_text(self.raw_text)  # Clean text before processing
#         self.nlp_name = spacy.load("en_core_web_lg")  # Better for extracting names
#         self.nlp_location = spacy.load("en_core_web_trf")  # Better for location
#         self.doc_name = self.nlp_name(self.text)
#         self.doc_location = self.nlp_location(self.text)


#     def extract_text(self, file_path):
#         """Extracts text from PDF or DOCX files."""
#         try:
#             if file_path.endswith(".pdf"):
#                 reader = PdfReader(file_path)
#                 return " ".join([page.extract_text() for page in reader.pages if page.extract_text()])
#             elif file_path.endswith(".docx"):
#                 doc = Document(file_path)
#                 return "\n".join([para.text for para in doc.paragraphs])
#             else:
#                 raise ValueError("Unsupported file format. Use PDF or DOCX.")
#         except Exception as e:
#             return f"Error extracting text: {e}"

#     def clean_text(self, text):
#         """Removes unwanted special characters while preserving emails, numbers, and dashes."""
#         return re.sub(r"[^\w\s@.+#&-]", "", text)

#     def extract_name(self):
#         """Extracts the person's name using the large NER model."""
#         persons = [ent.text for ent in self.doc_name.ents if ent.label_ == "PERSON"]
#         return persons[0] if persons else "Not found"

#     def extract_phone_numbers(self):
#         """Extracts phone numbers using `phonenumbers` library with a default country code."""
#         numbers = []
#         for match in phonenumbers.PhoneNumberMatcher(self.text, None):  # Default country: US
#             formatted_number = phonenumbers.format_number(match.number, phonenumbers.PhoneNumberFormat.INTERNATIONAL)
#             numbers.append(formatted_number)
#         return numbers if numbers else ""

#     def extract_email(self):
#         """Extracts an email address using regex."""
#         match = re.search(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}", self.text)
#         return match.group() if match else "Not found"

#     def extract_address(self):
#         for ent in self.doc_location.ents:
#             if ent.label_ == "GPE":
#                 return ent.text  # Return first detected location
#         return "Not found"

#     def extract_skills(self, skills_list):
#         """Matches keywords from a predefined skillset."""
#         return [skill for skill in skills_list if re.search(rf"\b{re.escape(skill)}\b", self.text, re.IGNORECASE)]

#     def extract_education(self):
#         """Extracts the education section from the resume text."""
#         EDUCATION_SECTIONS = [
#             "Education", "Academic History", "Qualifications", "Academic Background",
#             "Study History", "Education History", "Degrees", "ACADEMIC QUALIFICATION"
#         ]
       
#         NON_EDUCATION_SECTIONS = [
#             "Experience", "Work Experience", "Professional Experience", "Employment History",
#             "Job History", "Career History", "Projects", "Certifications", "Skills",
#             "Publications", "Research", "References"
#         ]
 
#         education_section = []
#         in_education_section = False
#         lines = self.text.split("\n")
 
#         for line in lines:
#             line = line.strip()
           
#             # Check if we are entering the education section
#             if any(re.fullmatch(rf"\b{section}\b", line, re.IGNORECASE) for section in EDUCATION_SECTIONS):
#                 in_education_section = True
#                 continue
 
#             # Stop if we reach another section (Experience, Projects, etc.)
#             if in_education_section and any(re.fullmatch(rf"\b{section}\b", line, re.IGNORECASE) for section in NON_EDUCATION_SECTIONS):
#                 break
 
#             # Capture education details
#             if in_education_section:
#                 education_section.append(line)
 
#         return "\n".join(education_section) if education_section else "No education section found"
 
 
#     def extract_experience(self):
#         """Extracts the experience section from the resume text."""
#         EXPERIENCE_SECTIONS = [
#             "Experience", "Work Experience", "Professional Experience", "Employment History",
#             "Job History", "Career History", "EMPLOYMENT DETAILS"
#         ]
       
#         NON_EXPERIENCE_SECTIONS = [
#             "Education", "Academic History", "Qualifications", "Academic Background",
#             "Study History", "Education History", "Degrees", "Projects", "Certifications",
#             "Skills", "Publications", "Research", "References"
#         ]
 
#         experience_section = []
#         in_experience_section = False
#         lines = self.text.split("\n")
 
#         for line in lines:
#             line = line.strip()
           
#             # Check if we are entering the experience section
#             if any(re.fullmatch(rf"\b{section}\b", line, re.IGNORECASE) for section in EXPERIENCE_SECTIONS):
#                 in_experience_section = True
#                 continue
 
#             # Stop if we reach another section (Education, Projects, etc.)
#             if in_experience_section and any(re.fullmatch(rf"\b{section}\b", line, re.IGNORECASE) for section in NON_EXPERIENCE_SECTIONS):
#                 break
 
#             # Capture experience details
#             if in_experience_section:
#                 experience_section.append(line)
 
#         return "\n".join(experience_section) if experience_section else "No experience section found"

# # Example Usage
# if __name__ == "__main__":
#     file_path = r"C:\Users\araul\Downloads\ATS classic HR resume.docx"
#     parser = ResumeParser(file_path)

import re
import phonenumbers
import spacy
from pypdf import PdfReader
from docx import Document
 
class ResumeParser:
    def __init__(self, file_path):
        self.file_path = file_path
        self.raw_text = self.extract_text(file_path)
        self.text = self.clean_text(self.raw_text)  # Clean text before processing
        self.nlp_name = spacy.load("en_core_web_lg")  # Better for extracting names
        self.nlp_location = spacy.load("en_core_web_trf")  # Better for location
        self.doc_name = self.nlp_name(self.text)
        self.doc_location = self.nlp_location(self.text)
 
 
    def extract_text(self, file_path):
        """Extracts text from PDF or DOCX files."""
        try:
            if file_path.endswith(".pdf"):
                reader = PdfReader(file_path)
                return " ".join([page.extract_text() for page in reader.pages if page.extract_text()])
            elif file_path.endswith(".docx"):
                doc = Document(file_path)
                return "\n".join([para.text for para in doc.paragraphs])
            else:
                raise ValueError("Unsupported file format. Use PDF or DOCX.")
        except Exception as e:
            return f"Error extracting text: {e}"
 
    def clean_text(self, text):
        """Removes unwanted special characters while preserving emails, numbers, and dashes."""
        return re.sub(r"[^\w\s@.+#&-]", "", text)
 
    def extract_name(self):
        """Extracts the person's name using the large NER model."""
        persons = [ent.text for ent in self.doc_name.ents if ent.label_ == "PERSON"]
        return persons[0] if persons else "Not found"
 
    def extract_phone_numbers(self):
        for match in phonenumbers.PhoneNumberMatcher(self.text, "US"):
            return phonenumbers.format_number(match.number, phonenumbers.PhoneNumberFormat.INTERNATIONAL)
        return ""
 
 
    def extract_email(self):
        """Extracts an email address using regex."""
        match = re.search(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}", self.text)
        return match.group() if match else "Not found"
 
    def extract_address(self):
        for ent in self.doc_location.ents:
            if ent.label_ == "GPE":
                return ent.text  # Return first detected location
        return "Not found"
 
 
    def extract_skills(self, skills_list):
        """Matches keywords from a predefined skillset."""
        return [skill for skill in skills_list if re.search(rf"\b{re.escape(skill)}\b", self.text, re.IGNORECASE)]
 
    def extract_education(self):
        """Extracts the education section from the resume text."""
        EDUCATION_SECTIONS = [
            "Education", "Academic History", "Qualifications", "Academic Background",
            "Study History", "Education History", "Degrees", "ACADEMIC QUALIFICATION"
        ]
       
        NON_EDUCATION_SECTIONS = [
            "Experience", "Work Experience", "Professional Experience", "Employment History",
            "Job History", "Career History", "Projects", "Certifications", "Skills",
            "Publications", "Research", "References"
        ]
 
        education_section = []
        in_education_section = False
        lines = self.text.split("\n")
 
        for line in lines:
            line = line.strip()
           
            # Check if we are entering the education section
            if any(re.fullmatch(rf"\b{section}\b", line, re.IGNORECASE) for section in EDUCATION_SECTIONS):
                in_education_section = True
                continue
 
            # Stop if we reach another section (Experience, Projects, etc.)
            if in_education_section and any(re.fullmatch(rf"\b{section}\b", line, re.IGNORECASE) for section in NON_EDUCATION_SECTIONS):
                break
 
            # Capture education details
            if in_education_section:
                education_section.append(line)
 
        return "\n".join(education_section) if education_section else "No education section found"
 
 
    def extract_experience(self):
        """Extracts the experience section from the resume text."""
        EXPERIENCE_SECTIONS = [
            "Experience", "Work Experience", "Professional Experience", "Employment History",
            "Job History", "Career History", "EMPLOYMENT DETAILS"
        ]
       
        NON_EXPERIENCE_SECTIONS = [
            "Education", "Academic History", "Qualifications", "Academic Background",
            "Study History", "Education History", "Degrees", "Projects", "Certifications",
            "Skills", "Publications", "Research", "References"
        ]
 
        experience_section = []
        in_experience_section = False
        lines = self.text.split("\n")
 
        for line in lines:
            line = line.strip()
           
            # Check if we are entering the experience section
            if any(re.fullmatch(rf"\b{section}\b", line, re.IGNORECASE) for section in EXPERIENCE_SECTIONS):
                in_experience_section = True
                continue
 
            # Stop if we reach another section (Education, Projects, etc.)
            if in_experience_section and any(re.fullmatch(rf"\b{section}\b", line, re.IGNORECASE) for section in NON_EXPERIENCE_SECTIONS):
                break
 
            # Capture experience details
            if in_experience_section:
                experience_section.append(line)
 
        return "\n".join(experience_section) if experience_section else "No experience section found"